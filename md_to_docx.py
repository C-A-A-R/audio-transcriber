# md_to_docx.py
import re
import tempfile
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt

BOLD_RE = re.compile(r"\*\*(.*?)\*\*", re.DOTALL)

def _apply_runs_with_bold(paragraph, text):
    """
    Divide el texto por **bold** markers y crea runs con bold donde corresponda.
    """
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            content = part[2:-2]
            run = paragraph.add_run(content)
            run.bold = True
        else:
            # normal text part (may contain inline formatting not supported)
            paragraph.add_run(part)

def markdown_to_docx(markdown_text: str) -> str:
    """
    Parsea un subset de Markdown y genera un .docx.
    Retorna la ruta al archivo .docx temporal.
    Soporta:
        - # heading
        - ## subheading
        - - bullets
        - 1. numbered lists
      - **bold**
    """
    doc = Document()
    # default font size
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = markdown_text.splitlines()
    i = 0
    in_list = False
    list_type = None  # 'bullet' or 'number'

    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            # blank line -> paragraph break
            doc.add_paragraph()
            i += 1
            continue

        # Heading level 1
        if line.startswith('# '):
            header = line[2:].strip()
            p = doc.add_heading(header.upper(), level=1)
            i += 1
            continue

        # Heading level 2
        if line.startswith('## '):
            header = line[3:].strip()
            p = doc.add_heading(header, level=2)
            i += 1
            continue

        # Numbered list like "1. something"
        m_num = re.match(r'^\s*\d+\.\s+(.*)', line)
        if m_num:
            content = m_num.group(1).strip()
            p = doc.add_paragraph(style='List Number')
            _apply_runs_with_bold(p, content)
            i += 1
            continue

        # Bullet list "- " or "• "
        m_bullet = re.match(r'^\s*[-\u2022]\s+(.*)', line)
        if m_bullet:
            content = m_bullet.group(1).strip()
            p = doc.add_paragraph(style='List Bullet')
            _apply_runs_with_bold(p, content)
            i += 1
            continue

        # Fallback: normal paragraph (may contain inline bold)
        p = doc.add_paragraph()
        _apply_runs_with_bold(p, line)
        i += 1

    # guardar docx en archivo temporal
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_path = tmp.name
    tmp.close()
    doc.save(tmp_path)
    return tmp_path
